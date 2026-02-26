#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Wrapper для lassio_transform.py который выводит результат в JSON формате
для использования веб-интерфейсом.
"""

import sys
import json
import argparse
import pandas as pd
from pathlib import Path
from typing import Any, List, Dict

# Попытка импорта python-docx для поддержки docx файлов
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Добавить папку Difference в путь для импорта
sys.path.insert(0, str(Path(__file__).parent.parent / 'Difference'))


def is_sku_code(text: str) -> bool:
    """
    Определяет, является ли текст SKU кодом товара.
    SKU коды имеют два паттерна:
    1. Полностью цифры (4 и более символов): например 12345
    2. Начинаются с "F0" и затем остальное цифры: например F0123, F01234
    """
    if not text or len(text) < 1:
        return False
    
    text = text.strip()
    
    # Пропускаем служебные слова
    if text.lower() in ['sku', 'код', 'kod', 'kod tovarа', 'nazvanie', 'kol-vo', 'итого', 'сумма', 'итоги']:
        return False
    
    # Паттерн 1: Полностью цифры (4 или более символов)
    if text.isdigit() and len(text) >= 4:
        return True
    
    # Паттерн 2: Начинается с "F0" и остальное цифры (F0XXXX)
    if text.startswith('F0') and len(text) >= 4:
        rest = text[2:]  # Всё после "F0"
        if rest.isdigit():
            return True
    
    return False


def extract_document_numbers_from_docx(file_path: str) -> List[str]:
    """
    Извлекает ВСЕ номера документов из Word файла.
    Ищет текст "Номер документа: " и берет цифры после него.
    """
    if not HAS_DOCX:
        return []
    
    document_numbers = []
    
    try:
        doc = Document(file_path)
        
        # Сначала ищем в параграфах
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if 'номер документа' in text.lower() or 'номер заказа' in text.lower():
                if ':' in text:
                    parts = text.split(':')
                    if len(parts) > 1:
                        number = parts[1].strip()
                        digits = ''.join(c for c in number if c.isdigit())
                        if digits and digits not in document_numbers:
                            document_numbers.append(digits)
                            print(f"📄 Word: Найден номер документа: {digits}", file=sys.stderr)
        
        # Если не нашли в параграфах - ищем в таблицах
        if not document_numbers:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if 'номер документа' in text.lower() or 'номер заказа' in text.lower():
                            if ':' in text:
                                parts = text.split(':')
                                if len(parts) > 1:
                                    number = parts[1].strip()
                                    digits = ''.join(c for c in number if c.isdigit())
                                    if digits and digits not in document_numbers:
                                        document_numbers.append(digits)
                                        print(f"📄 Word: Найден номер документа: {digits}", file=sys.stderr)
    except Exception as e:
        print(f"⚠️  Ошибка при поиске номеров документов в Word: {e}", file=sys.stderr)
    
    return document_numbers


def extract_document_numbers_from_excel(file_path: str, exord_mode: bool = False) -> dict:
    """
    Извлекает ВСЕ номера документов из Excel и привязывает товары к ним.
    Из значения "11813 - Статус IN_PROGRESS" берет только "11813"
    Возвращает словарь: { row_idx: "11813", ... }
    """
    try:
        df = pd.read_excel(file_path)
        df = df.dropna(how="all")
        df = df.reset_index(drop=True)
        
        # Ищем столбец с номерами документов - ищем более специфичное имя
        doc_col_name = None
        # Сначала ищем точное совпадение "Заказ №" или похожее
        for col_name in df.columns:
            col_str = str(col_name).lower().strip()
            if '№' in col_str or 'zakazno' in col_str.replace(' ', ''):
                doc_col_name = col_name
                print(f"📊 Excel: Найден столбец документов: {col_name}", file=sys.stderr)
                break
        
        # Если не нашли с №, ищем "Заказ" но НЕ "Заказано"
        if not doc_col_name:
            for col_name in df.columns:
                col_str = str(col_name).lower().strip()
                # Ищем "заказ" но исключаем "заказано"
                if 'заказ' in col_str and 'заказано' not in col_str:
                    doc_col_name = col_name
                    print(f"📊 Excel: Найден столбец документов: {col_name}", file=sys.stderr)
                    break
        
        # Если нашли столбец - привязываем каждого товара к его документу
        doc_mapping = {}  # row_idx -> doc_number
        
        if doc_col_name:
            for row_idx, val in enumerate(df[doc_col_name]):
                if pd.notna(val):
                    val_str = str(val).strip()
                    # Извлекаем ПЕРВОЕ число из строки (до пробела или другого символа)
                    # Например из "11813 - Статус IN_PROGRESS" берем "11813"
                    digits = ''
                    for char in val_str:
                        if char.isdigit():
                            digits += char
                        elif digits:
                            #停止если нашли число и потом нашли не-цифру
                            break
                    
                    if digits:
                        doc_mapping[row_idx] = digits
        
        # Возвращаем текущие номера
        unique_docs = sorted(set(doc_mapping.values())) if doc_mapping else []
        print(f"📊 Excel: Найдено {len(unique_docs)} уникальных номеров документов: {', '.join(unique_docs)}", file=sys.stderr)
        
        return doc_mapping, unique_docs
        
    except Exception as e:
        print(f"⚠️  Ошибка при поиске номеров документов в Excel: {e}", file=sys.stderr)
        return {}, []


def extract_from_docx(file_path: str, exord_mode: bool = False, exord_column: str = 'отправлено') -> Dict[str, List[Dict[str, Any]]]:
    """
    Извлекает данные из DOCX файла.
    Если exord_mode=True, ищет столбец exord_column ("Отправлено" или "Доставлено") и использует его значения.
    """
    if not HAS_DOCX:
        raise RuntimeError("Библиотека python-docx не установлена. Установите: pip install python-docx")
    
    doc = Document(file_path)
    if not doc.tables:
        raise ValueError("В документе нет таблиц")
    
    table = doc.tables[0]
    products: List[Dict[str, Any]] = []
    
    # Извлекаем номер документа
    doc_numbers = extract_document_numbers_from_docx(file_path)
    default_doc_number = doc_numbers[0] if doc_numbers else ""
    
    # Ищем столбцы в заголовке (первые несколько строк)
    sent_col_idx = -1
    delivered_col_idx = -1
    target_col_idx = -1

    sent_terms = ['отправлено', 'otpravleno', 'отправлен']
    delivered_terms = ['доставлено', 'dostavleno', 'доставлен']

    sent_candidates = []
    delivered_candidates = []

    for header_row_idx in range(min(5, len(table.rows))):
        header_row = table.rows[header_row_idx]
        for cell_idx, cell in enumerate(header_row.cells):
            cell_text = cell.text.strip().lower()
            if any(term in cell_text for term in sent_terms):
                sent_candidates.append(cell_idx)
            if any(term in cell_text for term in delivered_terms):
                delivered_candidates.append(cell_idx)

    if sent_candidates:
        sent_col_idx = sent_candidates[0]
    if delivered_candidates:
        delivered_col_idx = delivered_candidates[-1]

    if exord_mode:
        if exord_column == 'доставлено':
            target_col_idx = delivered_col_idx
        else:
            target_col_idx = sent_col_idx

        if target_col_idx >= 0:
            print(f"🔄 ЭКЗОРД: Найден столбец '{exord_column.upper()}' в позиции {target_col_idx}", file=sys.stderr)
        
        # Проверяем, есть ли числа в столбце
        has_numbers_in_target = False
        if target_col_idx >= 0:
            for check_row_idx in range(13, len(table.rows)):
                check_row = table.rows[check_row_idx]
                if target_col_idx < len(check_row.cells):
                    cell_text = check_row.cells[target_col_idx].text.strip()
                    if cell_text and cell_text not in ['-', '—', '–', '']:
                        try:
                            float(cell_text.replace(',', '.'))
                            has_numbers_in_target = True
                            print(f"   ✅ В столбце '{exord_column.upper()}' найдены числа", file=sys.stderr)
                            break
                        except:
                            pass
            
            if not has_numbers_in_target:
                print(f"   ⚠️ Столбец '{exord_column.upper()}' пуст или содержит только '-', используем первое число", file=sys.stderr)
                target_col_idx = -1  # Отключаем использование этого столбца
    
    # Обработка: начинаем с 14-й строки (индекс 13)
    for row_idx in range(13, len(table.rows)):
        row = table.rows[row_idx]
        
        # Проверяем служебные строки
        first_cell = row.cells[0].text.strip()
        if first_cell in ['Итого', 'Сумма', 'ИТОГО'] or not first_cell:
            continue
        
        if first_cell.lower() in ['sku', 'kod', 'kod tovarа', 'nazvanie', 'kol-vo', 'итоги']:
            continue
        
        # 1. Ищем SKU (код товара) в любой ячейке строки
        code = ""
        code_cell_idx = -1
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            if is_sku_code(cell_text):
                code = cell_text
                code_cell_idx = cell_idx
                break
        
        if not code:
            continue
        
        # 2. Название берем из той же строки (обычно после SKU)
        name = ""
        for cell_idx in range(code_cell_idx + 1, len(row.cells)):
            cell_text = row.cells[cell_idx].text.strip()
            
            if not cell_text:
                continue
            
            # Пропускаем числа
            if cell_text.replace('.', '').replace(',', '').isdigit():
                continue
            
            # Пропускаем если это похоже на SKU
            if is_sku_code(cell_text):
                continue
            
            name = cell_text
            break
        
        if not name:
            continue
        
        # 3. Количество
        quantity = None
        name_cell_idx = -1
        for cell_idx in range(code_cell_idx, len(row.cells)):
            if row.cells[cell_idx].text.strip() == name:
                name_cell_idx = cell_idx
                break
        
        if name_cell_idx >= 0:
            # Режим Экзорд: используем выбранный столбец если он найден и содержит числа
            if exord_mode and target_col_idx >= 0 and target_col_idx < len(row.cells):
                cell_text = row.cells[target_col_idx].text.strip()
                if cell_text and cell_text not in ['-', '—', '–']:
                    try:
                        qty_val = float(cell_text.replace(',', '.'))
                        if qty_val >= 0:
                            quantity = qty_val
                            print(f"   🔄 ЭКЗОРД: Взято из столбца '{exord_column.upper()}': {quantity}", file=sys.stderr)
                    except:
                        pass
            
            # Если не нашли - ищем первое число после названия
            if quantity is None:
                for cell_idx in range(name_cell_idx + 1, len(row.cells)):
                    cell_text = row.cells[cell_idx].text.strip()
                    if cell_text and cell_text not in ['-', '—', '–']:
                        try:
                            qty_val = float(cell_text.replace(',', '.'))
                            if qty_val >= 0:
                                quantity = qty_val
                                break
                        except:
                            continue
        
        # Режим Экзорд: ПРОПУСКАЕМ товар если выбранный столбец пустой или "-"
        if exord_mode and target_col_idx >= 0:
            if exord_column == 'доставлено' and sent_col_idx >= 0 and sent_col_idx < len(row.cells):
                sent_text = row.cells[sent_col_idx].text.strip()
                if not sent_text or sent_text in ['-', '—', '–']:
                    print(f"   ⏭️  ЭКЗОРД: Пропущен товар (нет ОТПРАВЛЕНО): '{name}'", file=sys.stderr)
                    continue
            else:
                cell_text = ""
                if target_col_idx < len(row.cells):
                    cell_text = row.cells[target_col_idx].text.strip()
                
                # Если в столбце нет числа - пропускаем этот товар
                if not cell_text or cell_text in ['-', '—', '–']:
                    print(f"   ⏭️  ЭКЗОРД: Пропущен товар (не {exord_column}): '{name}'", file=sys.stderr)
                    continue
        
        sent_qty = None
        delivered_qty = None

        if sent_col_idx >= 0 and sent_col_idx < len(row.cells):
            sent_text = row.cells[sent_col_idx].text.strip()
            if sent_text and sent_text not in ['-', '—', '–']:
                try:
                    sent_qty = float(sent_text.replace(',', '.'))
                except:
                    sent_qty = None

        if delivered_col_idx >= 0 and delivered_col_idx < len(row.cells):
            delivered_text = row.cells[delivered_col_idx].text.strip()
            if delivered_text and delivered_text not in ['-', '—', '–']:
                try:
                    delivered_qty = float(delivered_text.replace(',', '.'))
                except:
                    delivered_qty = None

        if exord_mode and exord_column == 'доставлено' and quantity is None:
            quantity = sent_qty if sent_qty is not None else 0

        products.append({
            'sku': code,
            'name': name,
            'qtn_invoice': quantity,
            'qtn_fact': quantity,
            'qtn_sent': sent_qty,
            'qtn_delivered': delivered_qty,
            'document_number': default_doc_number
        })
    
    return products


def extract_from_excel(file_path: str, exord_mode: bool = False, exord_column: str = 'отправлено') -> List[Dict[str, Any]]:
    """
    Извлекает данные из Excel файла (любой формат: .xlsx, .xls, .xlsm).
    Автоматически находит SKU, название и количество в каждой строке.
    Если exord_mode=True, ищет столбец exord_column ("Отправлено" или "Доставлено") и использует его значения.
    Также извлекает номер документа и привязывает его к каждому товару.
    """
    try:
        df = pd.read_excel(file_path)  # type: ignore
        
        # Удаляем полностью пустые строки
        df = df.dropna(how="all")  # type: ignore
        
        # Очищаем от служебных строк вверху (если есть)
        df = df.reset_index(drop=True)
        
        # Извлекаем номера документов и привязку к товарам
        doc_mapping, unique_docs = extract_document_numbers_from_excel(file_path, exord_mode)
        
        # Ищем столбцы в заголовке
        sent_col_name = None
        delivered_col_name = None
        target_col_name = None
        sent_terms = ['отправлено', 'otpravleno', 'отправлен']
        delivered_terms = ['доставлено', 'dostavleno', 'доставлен']

        sent_candidates = []
        delivered_candidates = []

        for col_name in df.columns:
            col_str = str(col_name).lower().strip()
            if any(term in col_str for term in sent_terms):
                sent_candidates.append(col_name)
            if any(term in col_str for term in delivered_terms):
                delivered_candidates.append(col_name)

        if sent_candidates:
            sent_col_name = sent_candidates[0]
        if delivered_candidates:
            delivered_col_name = delivered_candidates[-1]

        if exord_mode:
            target_col_name = delivered_col_name if exord_column == 'доставлено' else sent_col_name
            if target_col_name:
                print(f"🔄 ЭКЗОРД: Найден столбец '{exord_column.upper()}': {target_col_name}", file=sys.stderr)
            
            # Проверяем, есть ли числа в этом столбце
            if target_col_name:
                has_numbers = False
                for val in df[target_col_name]:
                    if pd.notna(val):
                        val_str = str(val).strip()
                        if val_str and val_str not in ['-', '—', '–']:
                            try:
                                float(val_str.replace(',', '.'))
                                has_numbers = True
                                print(f"   ✅ В столбце '{exord_column.upper()}' найдены числа", file=sys.stderr)
                                break
                            except:
                                pass
                
                if not has_numbers:
                    print(f"   ⚠️ Столбец '{exord_column.upper()}' пуст или содержит только '-', используем первое число", file=sys.stderr)
                    target_col_name = None
        
        products: List[Dict[str, Any]] = []
        
        print(f"\n🔍 Начинаю обработку Excel файла. Всего строк: {len(df)}", file=sys.stderr)
        
        for row_idx, row in df.iterrows():
            # Конвертируем row в dict для удобства
            row_data = row.to_dict()
            
            # Отладка: показываем первые 3 строки
            if row_idx < 3:
                print(f"\n📋 Строка {row_idx}: {list(row_data.values())[:5]}", file=sys.stderr)
            
            # 1. Ищем SKU (код товара) в любой ячейке строки
            sku = ""
            sku_col_idx = -1
            
            for col_idx, (col_name, cell_value) in enumerate(row_data.items()):
                if pd.isna(cell_value):
                    continue
                
                cell_text = str(cell_value).strip()
                
                # Отладка: проверяем каждую ячейку в первых 3 строках
                if row_idx < 3 and cell_text:
                    is_sku = is_sku_code(cell_text)
                    print(f"   Колонка {col_idx}: '{cell_text}' -> is_sku={is_sku}", file=sys.stderr)
                
                if is_sku_code(cell_text):
                    sku = cell_text
                    sku_col_idx = col_idx
                    print(f"   ✅ Найден SKU: '{sku}' в колонке {sku_col_idx}", file=sys.stderr)
                    break
            
            if not sku:
                if row_idx < 3:
                    print(f"   ❌ SKU не найден в строке {row_idx}", file=sys.stderr)
                continue
            
            # 2. Название - следующий не-числовой столбец после SKU
            name = ""
            name_col_idx = -1
            
            columns_list = list(row_data.keys())
            for col_idx in range(sku_col_idx + 1, len(columns_list)):
                cell_value = row_data[columns_list[col_idx]]
                
                if pd.isna(cell_value):
                    continue
                
                cell_text = str(cell_value).strip()
                
                if not cell_text:
                    continue
                
                # Пропускаем числовые значения
                try:
                    float(cell_text.replace(',', '.'))
                    continue  # Это число, пропускаем
                except ValueError:
                    pass  # Это текст, проверяем дальше
                
                # Пропускаем если это похоже на SKU
                if is_sku_code(cell_text):
                    continue
                
                name = cell_text
                name_col_idx = col_idx
                print(f"   ✅ Найдено название: '{name}' в колонке {name_col_idx}", file=sys.stderr)
                break
            
            if not name:
                print(f"   ❌ Название не найдено для SKU '{sku}'", file=sys.stderr)
                continue
            
            # 3. Количество
            quantity = None
            
            if name_col_idx >= 0:
                columns_list = list(row_data.keys())
                
                # Режим Экзорд: используем выбранный столбец если он найден и содержит числа
                if exord_mode and target_col_name and target_col_name in row_data:
                    cell_value = row_data[target_col_name]
                    if pd.notna(cell_value):
                        cell_text = str(cell_value).strip()
                        if cell_text and cell_text not in ['-', '—', '–']:
                            try:
                                qty_val = float(cell_text.replace(',', '.'))
                                if qty_val >= 0:
                                    quantity = qty_val
                                    print(f"   🔄 ЭКЗОРД: Взято из столбца '{exord_column.upper()}': {quantity}", file=sys.stderr)
                            except:
                                pass
                
                # Если не нашли - ищем первое число после названия
                if quantity is None:
                    for col_idx in range(name_col_idx + 1, len(columns_list)):
                        cell_value = row_data[columns_list[col_idx]]
                        
                        if pd.isna(cell_value):
                            continue
                        
                        cell_text = str(cell_value).strip()
                        
                        if not cell_text or cell_text in ['-', '—', '–']:
                            continue
                        
                        try:
                            qty_val = float(cell_text.replace(',', '.'))
                            if qty_val >= 0:
                                quantity = qty_val
                                break
                        except ValueError:
                            continue
            
            # Режим Экзорд: ПРОПУСКАЕМ товар если выбранный столбец пустой или "-"
            if exord_mode and target_col_name:
                if exord_column == 'доставлено' and sent_col_name and sent_col_name in row_data:
                    sent_val = row_data.get(sent_col_name)
                    sent_text = ""
                    if pd.notna(sent_val):
                        sent_text = str(sent_val).strip()
                    if not sent_text or sent_text in ['-', '—', '–']:
                        print(f"   ⏭️  ЭКЗОРД: Пропущен товар (нет ОТПРАВЛЕНО): '{name}'", file=sys.stderr)
                        continue
                else:
                    cell_value = row_data.get(target_col_name)
                    cell_text = ""
                    if pd.notna(cell_value):
                        cell_text = str(cell_value).strip()
                    
                    # Если в столбце нет числа - пропускаем этот товар
                    if not cell_text or cell_text in ['-', '—', '–']:
                        print(f"   ⏭️  ЭКЗОРД: Пропущен товар (не {exord_column}): '{name}'", file=sys.stderr)
                        continue
            
            # Добавляем товар если количество найдено
            if quantity is not None:
                # Определяем номер документа для этого товара
                doc_number = doc_mapping.get(row_idx, "")
                if not doc_number and unique_docs:
                    doc_number = unique_docs[0]  # Используем первый найденный номер если не привязан
                
                sent_qty = None
                delivered_qty = None

                if sent_col_name and sent_col_name in row_data:
                    sent_val = row_data[sent_col_name]
                    if pd.notna(sent_val):
                        sent_text = str(sent_val).strip()
                        if sent_text and sent_text not in ['-', '—', '–']:
                            try:
                                sent_qty = float(sent_text.replace(',', '.'))
                            except:
                                sent_qty = None

                if delivered_col_name and delivered_col_name in row_data:
                    delivered_val = row_data[delivered_col_name]
                    if pd.notna(delivered_val):
                        delivered_text = str(delivered_val).strip()
                        if delivered_text and delivered_text not in ['-', '—', '–']:
                            try:
                                delivered_qty = float(delivered_text.replace(',', '.'))
                            except:
                                delivered_qty = None

                print(f"   ✅ Добавлен товар: SKU='{sku}', Название='{name}', Количество={quantity}" + (f", Документ='{doc_number}'" if doc_number else ""), file=sys.stderr)
                if exord_mode and exord_column == 'доставлено' and quantity is None:
                    quantity = sent_qty if sent_qty is not None else 0

                products.append({
                    'sku': sku,
                    'name': name,
                    'qtn_invoice': quantity,
                    'qtn_fact': quantity,
                    'qtn_sent': sent_qty,
                    'qtn_delivered': delivered_qty,
                    'document_number': doc_number
                })
            else:
                print(f"   ❌ Количество не найдено для SKU '{sku}' (название: '{name}')", file=sys.stderr)
        
        print(f"\n✅ Обработка завершена. Найдено товаров: {len(products)}\n", file=sys.stderr)
        return products
        
    except Exception as e:
        raise RuntimeError(f"Ошибка при обработке Excel файла: {str(e)}")


def transform_to_json(file_path: str, exord_mode: bool = False, exord_column: str = 'отправлено') -> Dict[str, Any]:
    """
    Трансформирует файл Лассио в JSON формат для фронта.
    Автоматически определяет тип файла (XLSX или DOCX).
    Извлекает товары и группирует их по номерам документов.
    Возвращает структуру: { "document_groups": { "11813": [...], "11814": [...] } }
    """
    file_ext = Path(file_path).suffix.lower()
    
    # Извлекаем товары (они уже содержат document_number)
    if file_ext == '.docx':
        products = extract_from_docx(file_path, exord_mode, exord_column)
    else:
        # Для всех Excel файлов (.xlsx, .xls, .xlsm)
        products = extract_from_excel(file_path, exord_mode, exord_column)
    
    # Группируем товары по документам
    document_groups = {}
    for product in products:
        doc_num = product.get('document_number', 'unknown')
        if doc_num not in document_groups:
            document_groups[doc_num] = []
        document_groups[doc_num].append(product)
    
    # Логируем результат
    unique_docs = list(document_groups.keys())
    print(f"📊 Группировка товаров: {len(unique_docs)} документов", file=sys.stderr)
    for doc_num in sorted(unique_docs):
        print(f"   📋 Документ {doc_num}: {len(document_groups[doc_num])} товаров", file=sys.stderr)
    
    return {
        'document_groups': document_groups,
        'all_products': products,  # Для обратной совместимости
        'document_numbers': sorted(unique_docs)
    }


def main():
    parser = argparse.ArgumentParser(description='Transform Lassio Excel to JSON')
    parser.add_argument('--source', type=str, required=True, help='Path to source Excel file')
    parser.add_argument('--output', type=str, default=None, help='Output JSON file path (optional)')
    parser.add_argument('--exord', action='store_true', help='Режим Экзорд: фильтровать товары по столбцу')
    parser.add_argument('--exord-column', type=str, default='отправлено', help='Столбец для режима Экзорд ("отправлено" или "доставлено")')
    
    args = parser.parse_args()
    
    file_path = Path(args.source)
    
    if not file_path.exists():
        print(json.dumps({'error': f'Файл не найден: {args.source}'}), file=sys.stdout)
        sys.exit(1)
    
    try:
        print(f"\n🔵 Режим Экзорд: {'ВКЛЮЧЕН' if args.exord else 'ВЫКЛЮЧЕН'}", file=sys.stderr)
        if args.exord:
            print(f"📋 Выбранный столбец: {args.exord_column.upper()}\n", file=sys.stderr)
        result = transform_to_json(str(file_path), exord_mode=args.exord, exord_column=args.exord_column)
        
        # Выводим JSON в stdout
        print(json.dumps(result, ensure_ascii=False, indent=2))
        
        # Опционально сохраняем в файл
        if args.output:
            output_path = Path(args.output)
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
        
        sys.exit(0)
        
    except Exception as e:
        print(json.dumps({'error': str(e)}, ensure_ascii=False))
        sys.exit(1)


if __name__ == '__main__':
    main()
